#!/usr/bin/env python3
"""
Demo Document Generator - Creates a Word document with all 17 chart types
Using the FRONTEND_GUIDELINES.md design system
"""

import matplotlib
matplotlib.use('Agg')  # Non-interactive backend
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Design System Colors
# Primary colors
PRIMARY_COLOR = '#000000'      # Black
PRIMARY_LIGHT = '#333333'      # Primary light for large areas
SECONDARY_COLOR = '#E98532'    # Orange
SECONDARY_DARK = '#C97028'     # Darker Orange
SECONDARY_LIGHT = '#F9A552'    # Lighter Orange

# Semantic colors (updated for large-area fills)
SUCCESS_COLOR = '#C7FFBC'      # Soft green
ERROR_COLOR = '#FCD2CE'        # Soft red
INFO_COLOR = '#BEE7FF'         # Soft blue
WARNING_COLOR = '#FFF2CB'      # Soft yellow

TERTIARY_COLOR = '#6C757D'     # Gray

# Chart color palettes
# For large-area fills (bar charts, area charts, pie charts) - use light colors
LARGE_AREA_COLORS = [
    PRIMARY_LIGHT, SECONDARY_LIGHT, SUCCESS_COLOR, ERROR_COLOR, INFO_COLOR, WARNING_COLOR,
    SECONDARY_COLOR, SECONDARY_DARK, TERTIARY_COLOR
]

# For small-area fills (line charts, scatter plots, markers) - use normal colors
SMALL_AREA_COLORS = [
    PRIMARY_COLOR, SECONDARY_COLOR, SECONDARY_DARK, TERTIARY_COLOR,
    '#8B5CF6', '#EC4899', '#14B8A6', '#6366F1'
]

CHART_COLORS = SMALL_AREA_COLORS  # Default palette

# Setup matplotlib with design system
plt.rcParams.update({
    'font.family': 'serif',
    'font.serif': ['Times New Roman', 'DejaVu Serif'],
    'font.size': 18,
    'axes.titlesize': 24,
    'axes.labelsize': 18,
    'xtick.labelsize': 14,
    'ytick.labelsize': 14,
    'legend.fontsize': 14,
    'figure.figsize': (10, 6),
    'figure.dpi': 150,
    'savefig.dpi': 150,
    'savefig.bbox': 'tight',
    'axes.edgecolor': PRIMARY_COLOR,
    'axes.linewidth': 1,
    'grid.color': '#D0D0D0',
})

def create_chart_dir():
    """Create directory for chart images"""
    chart_dir = 'demo_charts'
    os.makedirs(chart_dir, exist_ok=True)
    return chart_dir

def save_chart(filename):
    """Save current figure and close"""
    plt.savefig(filename, dpi=150, bbox_inches='tight')
    plt.close()

# Chart 1: Combination Bar and Line Chart
def create_combination_chart(chart_dir):
    fig, ax1 = plt.subplots(figsize=(10, 6))

    categories = ['Q1', 'Q2', 'Q3', 'Q4']
    bar_data = [150, 230, 180, 320]
    line_data = [120, 190, 240, 280]

    ax1.bar(categories, bar_data, color=[PRIMARY_COLOR, SECONDARY_COLOR, PRIMARY_COLOR, SECONDARY_COLOR])
    ax1.set_ylabel('Revenue ($K)', fontsize=14)

    ax2 = ax1.twinx()
    ax2.plot(categories, line_data, color=SECONDARY_COLOR, marker='o', linewidth=2, markersize=8)
    ax2.set_ylabel('Customers', fontsize=14)

    plt.title('Combination Bar and Line Chart', fontsize=24, fontweight='bold')
    save_chart(f'{chart_dir}/combination_chart.png')

# Chart 2: 100% Stacked Bar Chart
def create_100_stacked_bar(chart_dir):
    categories = ['Product A', 'Product B', 'Product C', 'Product D']
    segment1 = [40, 25, 55, 30]
    segment2 = [35, 45, 25, 40]
    segment3 = [25, 30, 20, 30]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.barh(categories, segment1, label='Segment 1', color=PRIMARY_COLOR)
    ax.barh(categories, segment2, left=segment1, label='Segment 2', color=SECONDARY_COLOR)
    ax.barh(categories, segment3, left=np.add(segment1, segment2), label='Segment 3', color=TERTIARY_COLOR)

    plt.title('100% Stacked Bar Chart', fontsize=24, fontweight='bold')
    plt.legend(loc='lower right')
    save_chart(f'{chart_dir}/100_stacked_bar.png')

# Chart 3: Line Chart
def create_line_chart(chart_dir):
    months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun']
    data1 = [10, 25, 30, 45, 40, 60]
    data2 = [15, 20, 35, 30, 50, 55]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(months, data1, color=PRIMARY_COLOR, marker='o', linewidth=2, label='Series 1')
    ax.plot(months, data2, color=SECONDARY_COLOR, marker='s', linewidth=2, label='Series 2')

    plt.title('Line Chart', fontsize=24, fontweight='bold')
    plt.legend()
    plt.grid(True, alpha=0.3)
    save_chart(f'{chart_dir}/line_chart.png')

# Chart 4: Pie Chart
def create_pie_chart(chart_dir):
    labels = ['Category A', 'Category B', 'Category C', 'Category D']
    sizes = [35, 25, 25, 15]
    colors = [PRIMARY_COLOR, SECONDARY_COLOR, TERTIARY_COLOR, SUCCESS_COLOR]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')

    plt.title('Pie Chart', fontsize=24, fontweight='bold')
    save_chart(f'{chart_dir}/pie_chart.png')

# Chart 5: Radar Chart
def create_radar_chart(chart_dir):
    categories = ['Speed', 'Quality', 'Cost', 'Reliability', 'Innovation']
    values1 = [8, 7, 6, 9, 7]
    values2 = [6, 8, 7, 7, 9]

    angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
    values1 += values1[:1]
    values2 += values2[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(10, 6), subplot_kw=dict(projection='polar'))
    ax.plot(angles, values1, 'o-', color=PRIMARY_COLOR, linewidth=2, label='Product A')
    ax.fill(angles, values1, alpha=0.25, color=PRIMARY_COLOR)
    ax.plot(angles, values2, 's-', color=SECONDARY_COLOR, linewidth=2, label='Product B')
    ax.fill(angles, values2, alpha=0.25, color=SECONDARY_COLOR)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories)
    plt.title('Radar Chart', fontsize=24, fontweight='bold')
    plt.legend(loc='upper right')
    save_chart(f'{chart_dir}/radar_chart.png')

# Chart 6: Mekko Chart
def create_mekko_chart(chart_dir):
    segments = [
        [40, 30, 30],
        [25, 50, 25],
        [60, 20, 20],
        [35, 35, 30]
    ]
    categories = ['Market A', 'Market B', 'Market C', 'Market D']
    segment_names = ['Low', 'Medium', 'High']
    widths = [30, 25, 20, 25]

    fig, ax = plt.subplots(figsize=(10, 6))

    left = 0
    for i, (segment, width) in enumerate(zip(segments, widths)):
        bottom = 0
        for j, value in enumerate(segment):
            ax.bar(left, value/100*100, width=width, bottom=bottom/100*100,
                   color=CHART_COLORS[j % len(CHART_COLORS)], edgecolor='white')
            bottom += value
        left += width

    plt.title('Mekko Chart', fontsize=24, fontweight='bold')
    plt.xlim(0, 100)
    plt.ylim(0, 100)
    save_chart(f'{chart_dir}/mekko_chart.png')

# Chart 7: Area Chart
def create_area_chart(chart_dir):
    x = range(1, 11)
    y1 = [10, 15, 20, 25, 22, 30, 35, 32, 40, 45]
    y2 = [5, 10, 8, 15, 12, 18, 20, 17, 22, 25]

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.fill_between(x, y1, color=PRIMARY_COLOR, alpha=0.7, label='Area 1')
    ax.fill_between(x, y2, color=SECONDARY_COLOR, alpha=0.7, label='Area 2')

    plt.title('Area Chart', fontsize=24, fontweight='bold')
    plt.legend()
    plt.grid(True, alpha=0.3)
    save_chart(f'{chart_dir}/area_chart.png')

# Chart 8: Scatter Plot
def create_scatter_plot(chart_dir):
    np.random.seed(42)
    x = np.random.randn(50) * 10 + 50
    y = np.random.randn(50) * 10 + 50

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.scatter(x, y, c=PRIMARY_COLOR, alpha=0.6, s=100, edgecolors=SECONDARY_COLOR, linewidth=2)

    plt.title('Scatter Plot', fontsize=24, fontweight='bold')
    plt.xlabel('X Value')
    plt.ylabel('Y Value')
    plt.grid(True, alpha=0.3)
    save_chart(f'{chart_dir}/scatter_plot.png')

# Chart 9: Heatmap
def create_heatmap(chart_dir):
    data = np.random.rand(10, 10) * 100

    fig, ax = plt.subplots(figsize=(10, 8))
    im = ax.imshow(data, cmap='Oranges', aspect='auto')

    plt.title('Heatmap', fontsize=24, fontweight='bold')
    plt.colorbar(im, label='Value')
    save_chart(f'{chart_dir}/heatmap.png')

# Chart 10: Histogram
def create_histogram(chart_dir):
    np.random.seed(42)
    data = np.random.randn(1000) * 10 + 50

    fig, ax = plt.subplots(figsize=(10, 6))
    ax.hist(data, bins=20, color=PRIMARY_COLOR, edgecolor='white', alpha=0.8)

    plt.title('Histogram', fontsize=24, fontweight='bold')
    plt.xlabel('Value')
    plt.ylabel('Frequency')
    save_chart(f'{chart_dir}/histogram.png')

# Chart 11: Box Plot
def create_box_plot(chart_dir):
    np.random.seed(42)
    data = [
        np.random.randn(100) * 5 + 50,
        np.random.randn(100) * 5 + 55,
        np.random.randn(100) * 5 + 45,
        np.random.randn(100) * 5 + 60
    ]
    labels = ['Q1', 'Q2', 'Q3', 'Q4']

    fig, ax = plt.subplots(figsize=(10, 6))
    bp = ax.boxplot(data, labels=labels, patch_artist=True)

    for patch, color in zip(bp['boxes'], [PRIMARY_COLOR, SECONDARY_COLOR, TERTIARY_COLOR, SUCCESS_COLOR]):
        patch.set_facecolor(color)
        patch.set_alpha(0.7)

    plt.title('Box Plot', fontsize=24, fontweight='bold')
    plt.ylabel('Value')
    save_chart(f'{chart_dir}/box_plot.png')

# Chart 12: Violin Plot
def create_violin_plot(chart_dir):
    np.random.seed(42)
    data = [
        np.random.randn(100) * 5 + 50,
        np.random.randn(100) * 5 + 55,
        np.random.randn(100) * 5 + 45,
        np.random.randn(100) * 5 + 60
    ]
    labels = ['Group A', 'Group B', 'Group C', 'Group D']

    fig, ax = plt.subplots(figsize=(10, 6))
    parts = ax.violinplot(data, showmeans=True, showmedians=True)

    for i, pc in enumerate(parts['bodies']):
        pc.set_facecolor(CHART_COLORS[i % len(CHART_COLORS)])
        pc.set_alpha(0.7)

    plt.title('Violin Plot', fontsize=24, fontweight='bold')
    plt.ylabel('Value')
    save_chart(f'{chart_dir}/violin_plot.png')

# Chart 13: Funnel Chart
def create_funnel_chart(chart_dir):
    stages = ['Awareness', 'Interest', 'Consideration', 'Intent', 'Purchase', 'Loyalty']
    values = [100, 80, 60, 40, 25, 10]

    fig, ax = plt.subplots(figsize=(10, 6))

    # Use secondary color shades for gradient effect
    funnel_colors = [SECONDARY_COLOR, SECONDARY_DARK, SECONDARY_LIGHT,
                     SECONDARY_COLOR, SECONDARY_DARK, SECONDARY_LIGHT]

    for i, (stage, value) in enumerate(zip(stages, values)):
        width = value
        x = (100 - width) / 2
        ax.barh(stage, width, left=x, height=0.6,
                color=funnel_colors[i])
        ax.text(50, i, f'{value}%', ha='center', va='center',
                fontweight='bold', fontsize=12)

    plt.title('Funnel Chart', fontsize=24, fontweight='bold')
    plt.xlim(0, 100)
    save_chart(f'{chart_dir}/funnel_chart.png')

# Chart 14: Gauge Chart
def create_gauge_chart(chart_dir):
    fig, ax = plt.subplots(figsize=(10, 6), subplot_kw=dict(projection='polar'))

    value = 75
    colors = [ERROR_COLOR, SECONDARY_COLOR, SUCCESS_COLOR]
    ranges = [(0, 33), (33, 66), (66, 100)]

    for i, (start, end) in enumerate(ranges):
        start_rad = np.radians(start * 1.8)
        end_rad = np.radians(end * 1.8)
        ax.bar(x=start_rad, height=0.3, width=end_rad - start_rad,
               bottom=0.5, color=colors[i], alpha=0.7)

    ax.bar(x=np.radians(value * 1.8), height=0.4, width=0.02,
           bottom=0.5, color=PRIMARY_COLOR)

    ax.set_ylim(0, 1)
    ax.set_yticks([])
    ax.set_xticks([])

    plt.title('Gauge Chart', fontsize=24, fontweight='bold')
    plt.figtext(0.5, 0.3, f'{value}%', ha='center', fontsize=48, fontweight='bold')
    save_chart(f'{chart_dir}/gauge_chart.png')

# Chart 15: Treemap
def create_treemap(chart_dir):
    import squarify

    sizes = [40, 25, 15, 10, 5, 5]
    labels = ['Category A', 'Category B', 'Category C', 'Category D', 'Category E', 'Category F']
    colors_treemap = CHART_COLORS[:6]

    fig, ax = plt.subplots(figsize=(10, 6))
    squarify.plot(sizes=sizes, label=labels, color=colors_treemap,
                  alpha=0.8, text_kwargs={'fontsize': 12, 'color': 'white', 'fontweight': 'bold'})

    plt.title('Treemap', fontsize=24, fontweight='bold')
    plt.axis('off')
    save_chart(f'{chart_dir}/treemap.png')

# Chart 16: Sunburst Chart
def create_sunburst_chart(chart_dir):
    # Create a sunburst using nested pie charts
    fig, ax = plt.subplots(figsize=(10, 8), subplot_kw=dict(aspect='equal'))

    # Outer ring
    sizes_outer = [30, 25, 20, 25]
    colors_outer = [PRIMARY_COLOR, SECONDARY_COLOR, TERTIARY_COLOR, SUCCESS_COLOR]
    wedges1, _ = ax.pie(sizes_outer, radius=1, colors=colors_outer,
                        startangle=90, frame=True)

    # Inner ring
    sizes_inner = [15, 15, 10, 10, 20, 30]
    colors_inner = [PRIMARY_COLOR, '#333333', SECONDARY_COLOR, '#E99542',
                    TERTIARY_COLOR, '#8C8C8C']
    wedges2, _ = ax.pie(sizes_inner, radius=0.6, colors=colors_inner,
                        startangle=90)

    # Center circle
    circle = plt.Circle((0, 0), 0.3, color='white')
    ax.add_artist(circle)

    plt.title('Sunburst Chart', fontsize=24, fontweight='bold', pad=20)
    save_chart(f'{chart_dir}/sunburst_chart.png')

# Chart 17: Waterfall Chart
def create_waterfall_chart(chart_dir):
    categories = ['Start', '+A', '+B', '-C', '+D', '-E', 'End']
    values = [100, 30, 25, -20, 35, -15, 0]

    # Calculate cumulative values
    cumulative = [100]
    for v in values[1:-1]:
        cumulative.append(cumulative[-1] + v)
    cumulative.append(cumulative[-1])

    fig, ax = plt.subplots(figsize=(10, 6))

    # Plot bars
    x_positions = range(len(categories))
    colors = ['#8C8C8C']  # Start
    for v in values[1:-1]:
        if v > 0:
            colors.append(SUCCESS_COLOR)
        else:
            colors.append(ERROR_COLOR)
    colors.append('#8C8C8C')  # End

    bottom_values = [0]
    for i, v in enumerate(values[1:], 1):
        if v > 0:
            bottom_values.append(cumulative[i-1])
        else:
            bottom_values.append(cumulative[i])

    for i, (cat, val, bottom, color) in enumerate(zip(categories, values, bottom_values, colors)):
        if val >= 0:
            ax.bar(i, val, bottom=bottom, color=color, alpha=0.8, edgecolor='black')
        else:
            ax.bar(i, abs(val), bottom=bottom, color=color, alpha=0.8, edgecolor='black')

    # Add value labels
    for i, (cat, cum) in enumerate(zip(categories, cumulative)):
        ax.text(i, cum + 5, str(cum), ha='center', fontweight='bold')

    plt.title('Waterfall Chart', fontsize=24, fontweight='bold')
    plt.xticks(x_positions, categories)
    save_chart(f'{chart_dir}/waterfall_chart.png')

def create_document(chart_dir):
    """Create the Word document with all charts"""
    doc = Document()

    # Title
    title = doc.add_heading('Demo Chart Collection', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_paragraph('This document demonstrates all 17 chart types using the custom design system with:')

    # Design system specs
    specs = doc.add_paragraph()
    specs.add_run('Primary Color: ').bold = True
    specs.add_run('#000000 (Black)\n')
    specs.add_run('Secondary Color: ').bold = True
    specs.add_run('#E98532 (Orange)\n')
    specs.add_run('Font: ').bold = True
    specs.add_run('Times New Roman\n')
    specs.add_run('Border Radius: ').bold = True
    specs.add_run('0px (Sharp corners)')

    doc.add_page_break()

    # Chart list
    charts = [
        ('Combination Bar and Line Chart', 'combination_chart.png'),
        ('100% Stacked Bar Chart', '100_stacked_bar.png'),
        ('Line Chart', 'line_chart.png'),
        ('Pie Chart', 'pie_chart.png'),
        ('Radar Chart', 'radar_chart.png'),
        ('Mekko Chart', 'mekko_chart.png'),
        ('Area Chart', 'area_chart.png'),
        ('Scatter Plot', 'scatter_plot.png'),
        ('Heatmap', 'heatmap.png'),
        ('Histogram', 'histogram.png'),
        ('Box Plot', 'box_plot.png'),
        ('Violin Plot', 'violin_plot.png'),
        ('Funnel Chart', 'funnel_chart.png'),
        ('Gauge Chart', 'gauge_chart.png'),
        ('Treemap', 'treemap.png'),
        ('Sunburst Chart', 'sunburst_chart.png'),
        ('Waterfall Chart', 'waterfall_chart.png'),
    ]

    for chart_title, filename in charts:
        heading = doc.add_heading(chart_title, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        heading.runs[0].font.bold = True
        heading.runs[0].font.size = Pt(24)

        # Add chart image
        try:
            doc.add_picture(f'{chart_dir}/{filename}', width=Inches(6.5))
        except FileNotFoundError:
            doc.add_paragraph(f'[Chart image: {filename}]')

        doc.add_paragraph()

    return doc

def main():
    print("Creating demo charts...")
    chart_dir = create_chart_dir()

    # Create all charts
    chart_functions = [
        create_combination_chart,
        create_100_stacked_bar,
        create_line_chart,
        create_pie_chart,
        create_radar_chart,
        create_mekko_chart,
        create_area_chart,
        create_scatter_plot,
        create_heatmap,
        create_histogram,
        create_box_plot,
        create_violin_plot,
        create_funnel_chart,
        create_gauge_chart,
        create_treemap,
        create_sunburst_chart,
        create_waterfall_chart,
    ]

    for i, func in enumerate(chart_functions, 1):
        print(f"  Creating chart {i}/17: {func.__name__}")
        func(chart_dir)

    print("\nCreating Word document...")
    doc = create_document(chart_dir)

    output_file = 'demo_charts_collection.docx'
    doc.save(output_file)
    print(f"\nDone! Document saved as: {output_file}")
    print(f"Charts saved in: {chart_dir}/")

if __name__ == '__main__':
    main()
